import hashlib
import re
from io import BytesIO
from typing import Any

from fastapi import UploadFile

from app.core.exceptions import AppError

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

TECHNICAL_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "go",
    "rust",
    "c++",
    "sql",
    "postgresql",
    "mongodb",
    "redis",
    "fastapi",
    "django",
    "flask",
    "react",
    "next.js",
    "node.js",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "tensorflow",
    "pytorch",
    "scikit-learn",
    "langchain",
    "langgraph",
    "git",
    "linux",
}

SOFT_SKILLS = {
    "communication",
    "leadership",
    "teamwork",
    "problem solving",
    "collaboration",
    "ownership",
    "adaptability",
    "mentoring",
}

LANGUAGES = {"english", "hindi", "spanish", "french", "german", "kannada", "tamil", "telugu", "marathi"}


class ResumeParser:
    async def extract_upload_text(self, file: UploadFile, content: bytes) -> str:
        if file.content_type == PDF_MIME:
            return self.extract_pdf_text(content)
        if file.content_type == DOCX_MIME:
            return self.extract_docx_text(content)
        raise AppError(f"Unsupported resume type: {file.content_type}", 415, "unsupported_resume_type")

    def extract_pdf_text(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise AppError("PDF parsing dependency pypdf is not installed", 500, "parser_dependency_missing") from exc
        reader = PdfReader(BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return self._normalize_text(text)

    def extract_docx_text(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise AppError("DOCX parsing dependency python-docx is not installed", 500, "parser_dependency_missing") from exc
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
        return self._normalize_text("\n".join(paragraphs + table_cells))

    def parse_structured_content(self, text: str) -> dict[str, Any]:
        lower = text.lower()
        technical = sorted(skill for skill in TECHNICAL_SKILLS if re.search(rf"(?<![\w.+-]){re.escape(skill)}(?![\w.+-])", lower))
        soft = sorted(skill for skill in SOFT_SKILLS if skill in lower)
        languages = sorted(language.title() for language in LANGUAGES if re.search(rf"\b{re.escape(language)}\b", lower))
        sections = self._sections(text)
        projects = self._bullet_section(sections, "projects")
        education = self._bullet_section(sections, "education")
        experience = self._bullet_section(sections, "experience")
        certificates = self._bullet_section(sections, "certifications") or self._bullet_section(sections, "certificates")
        achievements = [item["description"] for item in self._bullet_section(sections, "achievements")]
        return {
            "skills": sorted(set(technical + soft)),
            "projects": projects,
            "education": education,
            "experience": experience,
            "certificates": certificates,
            "achievements": achievements,
            "technical_skills": technical,
            "soft_skills": soft,
            "languages": languages,
            "raw_text_hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
            "parse_metadata": {
                "word_count": len(text.split()),
                "section_names": sorted(sections.keys()),
                "parser": "heuristic-section-parser",
            },
        }

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _sections(self, text: str) -> dict[str, str]:
        known = [
            "summary",
            "skills",
            "technical skills",
            "projects",
            "education",
            "experience",
            "work experience",
            "certifications",
            "certificates",
            "achievements",
            "languages",
        ]
        matches = list(re.finditer(r"(?im)^\s*([A-Z][A-Za-z /&-]{2,40})\s*:?\s*$", text))
        sections: dict[str, str] = {}
        for index, match in enumerate(matches):
            name = match.group(1).strip().lower()
            if name not in known:
                continue
            start = match.end()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            sections[name] = text[start:end].strip()
        return sections

    def _bullet_section(self, sections: dict[str, str], name: str) -> list[dict[str, str]]:
        body = sections.get(name, "")
        items = [line.strip(" -•\t") for line in body.splitlines() if line.strip(" -•\t")]
        return [{"description": item[:1000]} for item in items[:30]]
